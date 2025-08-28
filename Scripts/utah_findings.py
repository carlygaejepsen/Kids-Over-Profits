import pandas as pd
from docx import Document

df = pd.read_csv("utah_citations_8-24-2025.csv")
doc = Document()

# Get all inspection date columns and their matching findings columns
date_cols = [c for c in df.columns if "Inspection" in c and "Date" in c]
findings_cols = [c for c in df.columns if "Inspection" in c and "Findings" in c]

for _, row in df.iterrows():
    last_date = None
    last_findings = None

    # Iterate over date/findings pairs in order
    for date_col, findings_col in zip(date_cols, findings_cols):
        date_val = row[date_col]
        if pd.notna(date_val) and str(date_val).strip() and str(date_val).strip() != "None":
            last_date = date_val
            last_findings = row[findings_col]

    # Skip if no findings
    if not last_findings or str(last_findings).strip().lower() in ("none", "nan", ""):
        continue

    doc.add_paragraph(f"Name: {row['Name']}")
    doc.add_paragraph(f"Address: {row['Address']}")
    doc.add_paragraph(f"Date of last inspection: {last_date}")
    doc.add_paragraph(f"Results: {last_findings}")
    doc.add_paragraph("")

doc.save("utah_citations_report_with_findings.docx")