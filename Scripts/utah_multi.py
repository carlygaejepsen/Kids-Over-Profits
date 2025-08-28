import pandas as pd
from docx import Document

df = pd.read_csv("utah_citations_8-24-2025.csv")
doc = Document()

# Identify all inspection date/type/findings column triplets
triplets = []
for i, col in enumerate(df.columns):
    if "Inspection" in col and "Date" in col:
        date_col = col
        type_col = df.columns[i+1]
        findings_col = df.columns[i+2]
        triplets.append((date_col, type_col, findings_col))

for _, row in df.iterrows():
    violations = []
    for date_col, type_col, findings_col in triplets:
        findings_val = str(row[findings_col]).strip()
        if findings_val.lower() not in ("", "none", "nan"):
            violations.append({
                "date": pd.to_datetime(row[date_col], errors="coerce"),
                "type": row[type_col],
                "findings": findings_val
            })

    # Only include facilities with multiple violations
    if len(violations) > 1:
        # Sort newest first
        violations.sort(key=lambda v: (v["date"] is not pd.NaT, v["date"]), reverse=True)

        doc.add_paragraph(f"Name: {row['Name']}")
        doc.add_paragraph(f"Address: {row['Address']}")
        doc.add_paragraph("Violations:")
        for v in violations:
            date_str = v["date"].strftime("%Y-%m-%d") if pd.notna(v["date"]) else "Unknown date"
            doc.add_paragraph(f"  - Date: {date_str}")
            doc.add_paragraph(f"    Type: {v['type']}")
            doc.add_paragraph(f"    Findings: {v['findings']}")
        doc.add_paragraph("")  # blank line between facilities

doc.save("utah_multiple_violations_sorted.docx")